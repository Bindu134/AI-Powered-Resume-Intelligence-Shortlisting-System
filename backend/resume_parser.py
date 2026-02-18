"""
Resume Parser — extracts structured data from PDF, DOCX, TXT, and images using:
1. PyPDF2 for PDF text extraction
2. python-docx for DOCX extraction
3. pytesseract for image OCR
4. Anthropic Claude for intelligent parsing
"""

import io
import os
import re
import uuid
import json
import logging
from datetime import datetime

import anthropic
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

from models import CandidateProfile, Project

logger = logging.getLogger(__name__)

PARSE_PROMPT = """You are an expert HR analyst. Extract structured information from the resume text below.

Return ONLY valid JSON with this exact schema:
{
  "name": "Full Name or Unknown",
  "email": "email@example.com or null",
  "phone": "phone number or null",
  "skills": ["skill1", "skill2", ...],
  "years_experience": 0.0,
  "education": ["Degree, University, Year", ...],
  "projects": [
    {"name": "Project Name", "description": "brief description", "technologies": ["tech1"]}
  ],
  "work_history": ["Job Title at Company (Year-Year)", ...]
}

Rules:
- skills: list ALL technical skills, tools, languages, frameworks mentioned
- years_experience: total years of professional work experience (estimate from dates)
- Keep descriptions concise
- If info not found, use empty list [] or null

Resume text:
{text}"""


class ResumeParser:
    def __init__(self):
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if api_key:
            self.client = anthropic.Anthropic(api_key=api_key)
            self.use_llm = True
        else:
            self.client = None
            self.use_llm = False
            logger.warning("No ANTHROPIC_API_KEY found. Using rule-based parser.")

    # ─────────────────────────────────────────────
    # TEXT EXTRACTORS
    # ─────────────────────────────────────────────

    def _extract_text_from_pdf(self, content: bytes) -> str:
        """Extract raw text from PDF bytes."""
        if not HAS_PYPDF2:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)
        return "\n".join(pages)

    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract raw text from DOCX bytes."""
        if not HAS_DOCX:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        doc = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables inside the DOCX
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)

    def _extract_text_from_txt(self, content: bytes) -> str:
        """Extract raw text from TXT bytes."""
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1")  # fallback encoding

    def _extract_text_from_image(self, content: bytes) -> str:
        """Extract text from image using OCR (pytesseract)."""
        if not HAS_OCR:
            raise ImportError("pytesseract or Pillow not installed. Run: pip install pytesseract Pillow")
        image = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(image)
        return text

    def _extract_text(self, content: bytes, filename: str) -> str:
        """Route to correct extractor based on file extension."""
        ext = filename.lower().split(".")[-1]

        if ext == "pdf":
            return self._extract_text_from_pdf(content)
        elif ext == "docx":
            return self._extract_text_from_docx(content)
        elif ext == "txt":
            return self._extract_text_from_txt(content)
        elif ext in ("png", "jpg", "jpeg", "tiff", "bmp", "webp"):
            return self._extract_text_from_image(content)
        else:
            logger.warning(f"Unsupported file format: {ext}. Attempting PDF parse.")
            return self._extract_text_from_pdf(content)

    # ─────────────────────────────────────────────
    # PARSERS
    # ─────────────────────────────────────────────

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        retry=retry_if_exception_type(Exception),
    )
    def _call_llm(self, text: str) -> dict:
        """Call Claude to parse resume text."""
        prompt = PARSE_PROMPT.format(text=text[:8000])  # cap tokens
        message = self.client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=2048,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = message.content[0].text.strip()
        # Extract JSON even if surrounded by markdown
        match = re.search(r"\{.*\}", raw, re.DOTALL)
        if match:
            return json.loads(match.group())
        raise ValueError("No JSON found in LLM response")

    def _rule_based_parse(self, text: str) -> dict:
        """Fallback rule-based parser when no API key."""
        # Extract email
        email_match = re.search(r"[\w.+-]+@[\w-]+\.\w+", text)
        email = email_match.group() if email_match else None

        # Extract phone
        phone_match = re.search(r"[\+\(]?\d[\d\s\-\(\)]{8,}", text)
        phone = phone_match.group().strip() if phone_match else None

        # Extract name (first line that looks like a name)
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        name = lines[0] if lines else "Unknown"

        # Common tech skills keyword extraction
        SKILLS_VOCAB = {
            "python", "java", "javascript", "typescript", "react", "vue", "angular",
            "node", "django", "flask", "fastapi", "sql", "mysql", "postgresql",
            "mongodb", "redis", "docker", "kubernetes", "aws", "gcp", "azure",
            "git", "linux", "tensorflow", "pytorch", "scikit-learn", "pandas",
            "numpy", "c++", "c#", "golang", "rust", "html", "css", "graphql",
            "rest", "kafka", "spark", "hadoop", "tableau", "power bi", "excel",
            "machine learning", "deep learning", "nlp", "computer vision",
        }
        text_lower = text.lower()
        skills = [s for s in SKILLS_VOCAB if s in text_lower]

        # Estimate experience from year ranges
        years = re.findall(r"(20\d{2}|19\d{2})", text)
        years_int = sorted(set(int(y) for y in years))
        exp = 0.0
        if len(years_int) >= 2:
            exp = float(years_int[-1] - years_int[0])
            exp = min(exp, 30)  # sanity cap

        return {
            "name": name,
            "email": email,
            "phone": phone,
            "skills": skills,
            "years_experience": exp,
            "education": [],
            "projects": [],
            "work_history": [],
        }

    # ─────────────────────────────────────────────
    # MAIN ENTRY
    # ─────────────────────────────────────────────

    async def parse(self, content: bytes, filename: str) -> CandidateProfile:
        """Main entry: file bytes → CandidateProfile. Supports PDF, DOCX, TXT, images."""
        try:
            text = self._extract_text(content, filename)
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {e}")
            text = ""

        if not text.strip():
            logger.warning(f"No text extracted from {filename}, using filename as fallback.")
            text = f"Resume file: {filename}"

        # Parse structured data
        if self.use_llm:
            try:
                data = self._call_llm(text)
            except Exception as e:
                logger.warning(f"LLM parse failed, using rule-based: {e}")
                data = self._rule_based_parse(text)
        else:
            data = self._rule_based_parse(text)

        # Build projects
        projects = []
        for p in data.get("projects", []):
            if isinstance(p, dict):
                projects.append(Project(
                    name=p.get("name", ""),
                    description=p.get("description", ""),
                    technologies=p.get("technologies", []),
                ))

        return CandidateProfile(
            id=str(uuid.uuid4()),
            name=data.get("name", "Unknown"),
            email=data.get("email"),
            phone=data.get("phone"),
            skills=data.get("skills", []),
            years_experience=float(data.get("years_experience", 0)),
            education=data.get("education", []),
            projects=projects,
            work_history=data.get("work_history", []),
            raw_text=text[:2000],
            filename=filename,
            uploaded_at=datetime.utcnow().isoformat(),
        )
```

---

**What changed from your original:**

1. Added `python-docx`, `pytesseract`, and `Pillow` imports with safe try/except blocks
2. Added `_extract_text_from_docx()` — handles DOCX including tables inside the document
3. Added `_extract_text_from_txt()` — handles TXT with UTF-8 and latin-1 fallback encoding
4. Added `_extract_text_from_image()` — OCR support for PNG, JPG, JPEG, TIFF, BMP, WEBP
5. Added `_extract_text()` — a smart router that picks the right extractor based on file extension
6. Updated `parse()` — now calls `_extract_text()` instead of directly calling `_extract_text_from_pdf()`

Also update your `requirements.txt` by adding:
```
python-docx
pytesseract
Pillow
